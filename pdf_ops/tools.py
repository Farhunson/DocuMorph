import os
import shutil
import subprocess
from typing import List, Tuple
from PIL import Image
import fitz  # PyMuPDF
from PyPDF2 import PdfReader, PdfWriter
from docx import Document
from pdf2docx import Converter

UPLOADS = os.path.join(os.path.dirname(__file__), "..", "uploads")
OUTPUTS = os.path.join(os.path.dirname(__file__), "..", "outputs")
os.makedirs(UPLOADS, exist_ok=True)
os.makedirs(OUTPUTS, exist_ok=True)

def base_noext(path:str) -> str:
    return os.path.splitext(os.path.basename(path))[0]

def out_path(name:str) -> str:
    return os.path.join(OUTPUTS, name)

def has_binary(cmd: str) -> bool:
    return shutil.which(cmd) is not None

# ---------- Merge ----------
def merge_pdfs(paths: List[str]) -> str:
    writer = PdfWriter()
    for p in paths:
        reader = PdfReader(p)
        if reader.is_encrypted:
            raise RuntimeError(f"Encrypted file requires password: {os.path.basename(p)}")
        for page in reader.pages:
            writer.add_page(page)
    out = out_path("merged.pdf")
    with open(out, "wb") as f:
        writer.write(f)
    return out

# ---------- Split (each page into its own file) ----------
def split_pdf(path: str) -> List[str]:
    reader = PdfReader(path)
    outputs = []
    base = base_noext(path)
    for i, page in enumerate(reader.pages, start=1):
        writer = PdfWriter()
        writer.add_page(page)
        out = out_path(f"{base}_page_{i}.pdf")
        with open(out, "wb") as f:
            writer.write(f)
        outputs.append(out)
    return outputs

# ---------- Compress (Ghostscript if available, else PyMuPDF re-save) ----------
def compress_pdf(path: str, quality: str = "screen") -> str:
    # quality: screen|ebook|printer|prepress
    out = out_path(f"{base_noext(path)}_compressed.pdf")
    if has_binary("gs"):
        # Ghostscript command
        cmd = [
            "gs", "-sDEVICE=pdfwrite", "-dCompatibilityLevel=1.5",
            f"-dPDFSETTINGS=/{quality}", "-dNOPAUSE", "-dQUIET", "-dBATCH",
            f"-sOutputFile={out}", path
        ]
        subprocess.run(cmd, check=True)
        return out
    # fallback using PyMuPDF re-writing (milder compression)
    doc = fitz.open(path)
    doc.save(out, deflate=True, clean=True, garbage=3)
    doc.close()
    return out

# ---------- Protect / Unlock ----------
def protect_pdf(path: str, password: str) -> str:
    reader = PdfReader(path)
    writer = PdfWriter()
    for p in reader.pages:
        writer.add_page(p)
    writer.encrypt(password)
    out = out_path(f"{base_noext(path)}_protected.pdf")
    with open(out, "wb") as f:
        writer.write(f)
    return out

def unlock_pdf(path: str, password: str) -> str:
    reader = PdfReader(path)
    if reader.is_encrypted:
        if not reader.decrypt(password):
            raise RuntimeError("Incorrect password.")
    writer = PdfWriter()
    for p in reader.pages:
        writer.add_page(p)
    out = out_path(f"{base_noext(path)}_unlocked.pdf")
    with open(out, "wb") as f:
        writer.write(f)
    return out

# ---------- Rotate ----------
def rotate_pdf(path: str, angle: int = 90) -> str:
    reader = PdfReader(path)
    writer = PdfWriter()
    for p in reader.pages:
        p.rotate(angle)
        writer.add_page(p)
    out = out_path(f"{base_noext(path)}_rotated_{angle}.pdf")
    with open(out, "wb") as f:
        writer.write(f)
    return out

# ---------- Watermark (PDF watermark first page over all pages) ----------
def watermark_pdf(path: str, watermark_pdf_path: str) -> str:
    wm_reader = PdfReader(watermark_pdf_path)
    wm_page = wm_reader.pages[0]
    reader = PdfReader(path)
    writer = PdfWriter()
    for p in reader.pages:
        p.merge_page(wm_page)
        writer.add_page(p)
    out = out_path(f"{base_noext(path)}_watermarked.pdf")
    with open(out, "wb") as f:
        writer.write(f)
    return out

# ---------- Signature image (PNG/JPG) placed bottom-right ----------
def sign_pdf_with_image(path: str, image_path: str, scale: float = 0.25) -> str:
    doc = fitz.open(path)
    for page in doc:
        rect = page.rect
        img = fitz.Pixmap(image_path)
        # scale image to width fraction
        target_w = rect.width * scale
        ratio = target_w / img.width
        target_h = img.height * ratio
        # bottom-right margin
        x1 = rect.x1 - target_w - 36
        y1 = rect.y1 - target_h - 36
        x2 = x1 + target_w
        y2 = y1 + target_h
        page.insert_image(fitz.Rect(x1, y1, x2, y2), filename=image_path, keep_proportion=True)
    out = out_path(f"{base_noext(path)}_signed.pdf")
    doc.save(out)
    doc.close()
    return out

# ---------- Extract text ----------
def extract_text(path: str) -> str:
    out = out_path(f"{base_noext(path)}_text.txt")
    doc = fitz.open(path)
    with open(out, "w", encoding="utf-8") as f:
        for i, page in enumerate(doc, start=1):
            f.write(f"--- Page {i} ---\n")
            f.write(page.get_text())
            f.write("\n\n")
    doc.close()
    return out

# ---------- PDF → DOCX ----------
def pdf_to_docx(path: str) -> str:
    out = out_path(f"{base_noext(path)}_converted.docx")
    try:
        # best-effort layout conversion
        cv = Converter(path)
        cv.convert(out, start=0, end=None)
        cv.close()
        return out
    except Exception:
        # fallback: text-only
        doc = fitz.open(path)
        d = Document()
        for i, p in enumerate(doc, start=1):
            d.add_paragraph(f"--- Page {i} ---")
            d.add_paragraph(p.get_text())
        d.save(out)
        doc.close()
        return out

# ---------- PDF ↔ Images ----------
def pdf_to_images(path: str, fmt: str = "png") -> List[str]:
    fmt = fmt.lower()
    assert fmt in ("png", "jpg", "jpeg")
    doc = fitz.open(path)
    outs = []
    for i, page in enumerate(doc, start=1):
        pix = page.get_pixmap()
        fn = out_path(f"{base_noext(path)}_page_{i}.{ 'jpg' if fmt in ('jpg','jpeg') else 'png'}")
        pix.save(fn)
        outs.append(fn)
    doc.close()
    return outs

def images_to_pdf(image_paths: List[str]) -> str:
    imgs = [Image.open(p).convert("RGB") for p in image_paths]
    out = out_path("images_to_pdf.pdf")
    if not imgs:
        raise RuntimeError("No images provided.")
    first, rest = imgs[0], imgs[1:]
    first.save(out, save_all=True, append_images=rest)
    for im in imgs:
        im.close()
    return out

# ---------- Office → PDF (via LibreOffice) ----------
def office_to_pdf(path: str) -> str:
    """
    Convert DOC/DOCX/XLS/XLSX/PPT/PPTX to PDF using LibreOffice (soffice).
    """
    if not has_binary("soffice"):
        raise RuntimeError("LibreOffice (soffice) not found on PATH.")
    out_dir = OUTPUTS
    cmd = ["soffice", "--headless", "--convert-to", "pdf", "--outdir", out_dir, path]
    subprocess.run(cmd, check=True)
    # LibreOffice writes with original base name + .pdf
    out = out_path(f"{base_noext(path)}.pdf")
    if not os.path.isfile(out):
        raise RuntimeError("Conversion failed (LibreOffice).")
    return out

# ---------- Extract Images (all embedded images from PDF) ----------
def extract_images(path: str) -> List[str]:
    outs = []
    doc = fitz.open(path)
    for page_num, page in enumerate(doc, start=1):
        for img_index, img in enumerate(page.get_images(full=True), start=1):
            xref = img[0]
            pix = fitz.Pixmap(doc, xref)
            if pix.n > 4:  # CMYK
                pix = fitz.Pixmap(fitz.csRGB, pix)
            fn = out_path(f"{base_noext(path)}_p{page_num}_{img_index}.png")
            pix.save(fn)
            outs.append(fn)
            pix = None
    doc.close()
    return outs


# ---------- PDF → Excel (table extraction) ----------
def pdf_to_excel(path: str) -> str:
    """
    Extracts tables into XLSX. Requires camelot or tabula.
    """
    out = out_path(f"{base_noext(path)}.xlsx")
    try:
        import camelot
        tables = camelot.read_pdf(path, pages="all")
        if tables:
            tables.export(out, f="excel")
            return out
        else:
            raise RuntimeError("No tables found.")
    except Exception:
        # fallback: dump text into XLSX
        import openpyxl
        doc = fitz.open(path)
        wb = openpyxl.Workbook()
        ws = wb.active
        for i, page in enumerate(doc, start=1):
            ws.append([f"--- Page {i} ---"])
            ws.append([page.get_text()])
            ws.append([])
        wb.save(out)
        doc.close()
        return out


# ---------- PDF → HTML ----------
def pdf_to_html(path: str) -> str:
    out = out_path(f"{base_noext(path)}.html")
    doc = fitz.open(path)
    html = ["<html><body>"]
    for i, page in enumerate(doc, start=1):
        html.append(f"<h2>Page {i}</h2>")
        html.append("<pre>")
        html.append(page.get_text("text"))
        html.append("</pre>")
    html.append("</body></html>")
    with open(out, "w", encoding="utf-8") as f:
        f.write("\n".join(html))
    doc.close()
    return out


# ---------- OCR PDF ----------
def pdf_ocr(path: str, lang: str = "eng") -> str:
    """
    OCR scanned PDF into searchable PDF.
    Requires pytesseract and tesseract installed.
    """
    from pytesseract import image_to_pdf_or_hocr
    out = out_path(f"{base_noext(path)}_ocr.pdf")
    doc = fitz.open(path)
    pdf_bytes = b""
    for page in doc:
        pix = page.get_pixmap(dpi=300)
        img_bytes = pix.tobytes("png")
        pdf_bytes += image_to_pdf_or_hocr(Image.open(
            fitz.BytesIO(img_bytes)), lang=lang, extension="pdf")
    with open(out, "wb") as f:
        f.write(pdf_bytes)
    doc.close()
    return out


# ---------- Reorder Pages ----------
def reorder_pages(path: str, new_order: List[int]) -> str:
    reader = PdfReader(path)
    writer = PdfWriter()
    num_pages = len(reader.pages)
    for i in new_order:
        if 1 <= i <= num_pages:
            writer.add_page(reader.pages[i-1])
    out = out_path(f"{base_noext(path)}_reordered.pdf")
    with open(out, "wb") as f:
        writer.write(f)
    return out
